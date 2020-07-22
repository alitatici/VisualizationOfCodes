from docx import Document
from docx.shared import Inches

def eq():

    document = Document()
    p = document.add_paragraph("m = ({}  {}\n{}  {})".format)
    # run = p.add_run()

    document.save("123.docx")

    return document

eq()




