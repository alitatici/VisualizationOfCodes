from docx import Document
from docx.shared import Inches

document=Document()
document.add_heading("Structural Dynamic Reports")
document.add_paragraph("This is sDyna Tutorial")
a=2
b=a/2
b=int(b)
table = document.add_table(rows=a, cols=3)

cell = table.cell(b, 0)
cell.text = 'Mass(m)='
cell=table.cell(0,1)
cell.text='a'
cell=table.cell(0,2)
cell.text='a'
cell=table.cell(1,1)
cell.text='a'
cell=table.cell(1,2)
cell.text='a'


document.save("word1.docx")
